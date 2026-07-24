from __future__ import annotations

import argparse
import csv
import re
import time
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from sqlalchemy import text

from auction_etl.browser.manager import browser
from auction_etl.database.session import engine


JST = ZoneInfo("Asia/Tokyo")
UTC = ZoneInfo("UTC")
NEW_YORK = ZoneInfo("America/New_York")

DATE_FORMATS = (
    "%d %b %Y %H:%M:%S",
    "%d %B %Y %H:%M:%S",
    "%b %d, %Y %H:%M:%S",
    "%B %d, %Y %H:%M:%S",
    "%Y-%m-%d %H:%M:%S",
)

PATTERNS = {
    "auction_id": (
        re.compile(
            r"Auction\s+ID\s*[\r\n:]*\s*([A-Za-z0-9_-]+)",
            re.IGNORECASE,
        ),
    ),
    "opening_time": (
        re.compile(
            r"Opening\s+Time\s*\(JST\)\s*[\r\n:]*\s*"
            r"([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4}\s+"
            r"[0-9]{2}:[0-9]{2}:[0-9]{2})",
            re.IGNORECASE,
        ),
        re.compile(
            r"Opening\s+Time\s*\(JST\)\s*[\r\n:]*\s*"
            r"([A-Za-z]+\s+[0-9]{1,2},\s+[0-9]{4}\s+"
            r"[0-9]{2}:[0-9]{2}:[0-9]{2})",
            re.IGNORECASE,
        ),
    ),
    "closing_time": (
        re.compile(
            r"Closing\s+Time\s*\(JST\)\s*[\r\n:]*\s*"
            r"([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4}\s+"
            r"[0-9]{2}:[0-9]{2}:[0-9]{2})",
            re.IGNORECASE,
        ),
        re.compile(
            r"Closing\s+Time\s*\(JST\)\s*[\r\n:]*\s*"
            r"([A-Za-z]+\s+[0-9]{1,2},\s+[0-9]{4}\s+"
            r"[0-9]{2}:[0-9]{2}:[0-9]{2})",
            re.IGNORECASE,
        ),
    ),
    "starting_price": (
        re.compile(
            r"Starting\s+Price\s*[\r\n:]*\s*"
            r"([0-9][0-9,]*)\s*YEN",
            re.IGNORECASE,
        ),
    ),
    "current_price": (
        re.compile(
            r"Current\s+Price\s*[\r\n:]*\s*"
            r"([0-9][0-9,]*)\s*YEN",
            re.IGNORECASE,
        ),
    ),
    "bid_count": (
        re.compile(
            r"Number\s+of\s+Bids\s*[\r\n:]*\s*"
            r"([0-9][0-9,]*)",
            re.IGNORECASE,
        ),
    ),
    "seller": (
        re.compile(
            r"Seller\s*[\r\n:]+\s*([^\r\n]+)",
            re.IGNORECASE,
        ),
    ),
    "condition": (
        re.compile(
            r"Item\s+Condition\s*[\r\n:]+\s*([^\r\n]+)",
            re.IGNORECASE,
        ),
    ),
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Enrich Buyee auctions from detail pages."
    )
    parser.add_argument(
        "--profile",
        default="anonymous",
    )
    parser.add_argument(
        "--wait-seconds",
        type=float,
        default=2.0,
    )
    parser.add_argument(
        "--limit",
        type=int,
    )
    parser.add_argument(
        "--force",
        action="store_true",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("exports/buyee"),
    )
    return parser.parse_args()


def normalize_text(value: str) -> str:
    lines: list[str] = []

    for line in value.replace("\xa0", " ").splitlines():
        cleaned = " ".join(line.split())

        if cleaned:
            lines.append(cleaned)

    return "\n".join(lines)


def match_field(
    visible_text: str,
    field: str,
) -> str | None:
    for pattern in PATTERNS[field]:
        match = pattern.search(visible_text)

        if match is not None:
            return match.group(1).strip()

    return None


def parse_integer(value: str | None) -> int | None:
    if value is None:
        return None

    return int(value.replace(",", ""))


def parse_decimal(value: str | None) -> Decimal | None:
    if value is None:
        return None

    return Decimal(value.replace(",", ""))


def parse_jst_datetime(
    value: str | None,
) -> datetime | None:
    if value is None:
        return None

    for date_format in DATE_FORMATS:
        try:
            parsed = datetime.strptime(
                value,
                date_format,
            )
            return parsed.replace(tzinfo=JST)
        except ValueError:
            continue

    return None


def iso_in_zone(
    value: datetime | None,
    zone: ZoneInfo,
) -> str | None:
    if value is None:
        return None

    return value.astimezone(zone).isoformat()


def create_detail_table() -> None:
    statement = text(
        """
        CREATE TABLE IF NOT EXISTS warehouse.auction_detail (
            marketplace VARCHAR(32) NOT NULL,
            listing_id VARCHAR(128) NOT NULL,
            auction_id VARCHAR(128),
            started_at TIMESTAMPTZ,
            ended_at TIMESTAMPTZ,
            start_price NUMERIC(18, 2),
            detail_price NUMERIC(18, 2),
            bid_count INTEGER,
            seller TEXT,
            condition_text TEXT,
            detail_url TEXT,
            enriched_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            PRIMARY KEY (marketplace, listing_id)
        )
        """
    )

    with engine.begin() as connection:
        connection.execute(statement)


def load_candidates(
    force: bool,
    limit: int | None,
) -> list[dict[str, Any]]:
    conditions = [
        "a.marketplace = 'buyee'",
        "a.auction_url IS NOT NULL",
        "a.auction_url <> ''",
    ]

    if not force:
        conditions.append(
            """
            (
                d.listing_id IS NULL
                OR d.ended_at IS NULL
                OR d.start_price IS NULL
            )
            """
        )

    limit_sql = ""

    if limit is not None:
        limit_sql = "LIMIT :limit"

    statement = text(
        f"""
        SELECT
            a.id,
            a.listing_id,
            a.auction_url,
            a.title,
            a.seller,
            a.bid_count,
            a.start_price,
            a.final_price,
            a.tax_amount,
            a.gross_price,
            a.currency,
            a.ended_at
        FROM warehouse.auction AS a
        LEFT JOIN warehouse.auction_detail AS d
          ON d.marketplace = a.marketplace
         AND d.listing_id = a.listing_id
        WHERE {" AND ".join(conditions)}
        ORDER BY
            a.ended_at DESC NULLS LAST,
            a.id DESC
        {limit_sql}
        """
    )

    parameters: dict[str, Any] = {}

    if limit is not None:
        parameters["limit"] = limit

    with engine.connect() as connection:
        rows = connection.execute(
            statement,
            parameters,
        ).mappings().all()

    return [dict(row) for row in rows]


def extract_detail(
    page,
    row: dict[str, Any],
    wait_seconds: float,
) -> dict[str, Any]:
    response = page.goto(
        row["auction_url"],
        wait_until="domcontentloaded",
        timeout=120_000,
    )

    page.wait_for_timeout(
        max(1_000, int(wait_seconds * 1_000))
    )

    try:
        page.wait_for_selector(
            "text=Auction ID",
            timeout=20_000,
        )
    except Exception:
        pass

    visible_text = normalize_text(
        page.locator("body").inner_text(
            timeout=15_000
        )
    )

    listing_id = str(row["listing_id"])
    diagnostic_dir = Path("logs/buyee-detail")
    diagnostic_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    opening_text = match_field(
        visible_text,
        "opening_time",
    )
    closing_text = match_field(
        visible_text,
        "closing_time",
    )

    result = {
        "marketplace": "buyee",
        "listing_id": listing_id,
        "auction_id": match_field(
            visible_text,
            "auction_id",
        ),
        "started_at": parse_jst_datetime(
            opening_text
        ),
        "ended_at": parse_jst_datetime(
            closing_text
        ),
        "start_price": parse_decimal(
            match_field(
                visible_text,
                "starting_price",
            )
        ),
        "detail_price": parse_decimal(
            match_field(
                visible_text,
                "current_price",
            )
        ),
        "bid_count": parse_integer(
            match_field(
                visible_text,
                "bid_count",
            )
        ),
        "seller": match_field(
            visible_text,
            "seller",
        ),
        "condition_text": match_field(
            visible_text,
            "condition",
        ),
        "detail_url": page.url,
        "http_status": (
            response.status
            if response is not None
            else None
        ),
        "page_title": page.title(),
    }

    required = (
        result["auction_id"],
        result["ended_at"],
    )

    if not all(required):
        text_path = diagnostic_dir / f"{listing_id}.txt"
        html_path = diagnostic_dir / f"{listing_id}.html"
        screenshot_path = diagnostic_dir / f"{listing_id}.png"

        text_path.write_text(
            visible_text + "\n",
            encoding="utf-8",
        )
        html_path.write_text(
            page.content(),
            encoding="utf-8",
        )
        page.screenshot(
            path=str(screenshot_path),
            full_page=True,
        )

        raise RuntimeError(
            f"Required fields missing for {listing_id}. "
            f"Diagnostics: {text_path}"
        )

    return result


def save_detail(detail: dict[str, Any]) -> None:
    upsert_detail = text(
        """
        INSERT INTO warehouse.auction_detail (
            marketplace,
            listing_id,
            auction_id,
            started_at,
            ended_at,
            start_price,
            detail_price,
            bid_count,
            seller,
            condition_text,
            detail_url,
            enriched_at
        )
        VALUES (
            :marketplace,
            :listing_id,
            :auction_id,
            :started_at,
            :ended_at,
            :start_price,
            :detail_price,
            :bid_count,
            :seller,
            :condition_text,
            :detail_url,
            NOW()
        )
        ON CONFLICT (marketplace, listing_id)
        DO UPDATE SET
            auction_id = EXCLUDED.auction_id,
            started_at = EXCLUDED.started_at,
            ended_at = EXCLUDED.ended_at,
            start_price = EXCLUDED.start_price,
            detail_price = EXCLUDED.detail_price,
            bid_count = EXCLUDED.bid_count,
            seller = EXCLUDED.seller,
            condition_text = EXCLUDED.condition_text,
            detail_url = EXCLUDED.detail_url,
            enriched_at = NOW()
        """
    )

    update_auction = text(
        """
        UPDATE warehouse.auction
        SET
            ended_at = COALESCE(
                :ended_at,
                ended_at
            ),
            start_price = COALESCE(
                :start_price,
                start_price
            ),
            bid_count = COALESCE(
                :bid_count,
                bid_count
            ),
            seller = COALESCE(
                NULLIF(:seller, ''),
                seller
            )
        WHERE marketplace = :marketplace
          AND listing_id = :listing_id
        """
    )

    with engine.begin() as connection:
        connection.execute(
            upsert_detail,
            detail,
        )
        connection.execute(
            update_auction,
            detail,
        )


def export_rows() -> list[dict[str, Any]]:
    statement = text(
        """
        SELECT
            a.id,
            a.marketplace,
            a.listing_id,
            a.auction_url,
            a.seller,
            a.artist,
            a.title,
            a.media_type,
            a.disc_count,
            a.edition,
            a.catalog_number,
            a.condition_media,
            a.condition_cover,
            a.bulk_lot,
            a.bid_count,
            a.watch_count,
            a.start_price,
            a.final_price AS hammer_price_jpy,
            a.tax_rate,
            a.tax_amount AS tax_jpy,
            a.gross_price AS gross_price_jpy,
            a.price_includes_tax,
            a.currency,
            d.auction_id,
            d.started_at,
            a.ended_at,
            d.condition_text,
            d.detail_price AS detail_page_price_jpy,
            d.enriched_at
        FROM warehouse.auction AS a
        LEFT JOIN warehouse.auction_detail AS d
          ON d.marketplace = a.marketplace
         AND d.listing_id = a.listing_id
        WHERE a.marketplace = 'buyee'
        ORDER BY
            a.ended_at DESC NULLS LAST,
            a.id DESC
        """
    )

    with engine.connect() as connection:
        database_rows = connection.execute(
            statement
        ).mappings().all()

    rows: list[dict[str, Any]] = []

    for database_row in database_rows:
        row = dict(database_row)

        started_at = row.pop("started_at")
        ended_at = row.pop("ended_at")
        enriched_at = row.pop("enriched_at")

        row["started_at_jst"] = iso_in_zone(
            started_at,
            JST,
        )
        row["started_at_utc"] = iso_in_zone(
            started_at,
            UTC,
        )
        row["started_at_new_york"] = iso_in_zone(
            started_at,
            NEW_YORK,
        )

        row["ended_at_jst"] = iso_in_zone(
            ended_at,
            JST,
        )
        row["ended_at_utc"] = iso_in_zone(
            ended_at,
            UTC,
        )
        row["ended_at_new_york"] = iso_in_zone(
            ended_at,
            NEW_YORK,
        )

        row["enriched_at_utc"] = iso_in_zone(
            enriched_at,
            UTC,
        )

        rows.append(row)

    return rows


def export_csv(
    rows: list[dict[str, Any]],
    output_path: Path,
) -> Path:
    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    if not rows:
        raise RuntimeError(
            "No Buyee rows were available for export."
        )

    with output_path.open(
        "w",
        encoding="utf-8-sig",
        newline="",
    ) as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=list(rows[0]),
        )
        writer.writeheader()
        writer.writerows(rows)

    return output_path


def export_xlsx(
    rows: list[dict[str, Any]],
    output_path: Path,
) -> Path:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Buyee Auctions"

    headers = list(rows[0])
    worksheet.append(headers)

    for cell in worksheet[1]:
        cell.font = Font(bold=True)

    for row in rows:
        worksheet.append(
            [
                row.get(header)
                for header in headers
            ]
        )

    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions

    for index, header in enumerate(
        headers,
        start=1,
    ):
        values = [
            str(header),
            *[
                str(row.get(header) or "")
                for row in rows[:200]
            ],
        ]

        width = min(
            max(len(value) for value in values) + 2,
            60,
        )

        worksheet.column_dimensions[
            get_column_letter(index)
        ].width = width

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )
    workbook.save(output_path)

    return output_path


def export_docx(
    rows: list[dict[str, Any]],
    output_path: Path,
) -> Path:
    document = Document()
    document.add_heading(
        "Buyee Auctions — Detail-Enriched",
        level=0,
    )

    document.add_paragraph(
        f"Listings: {len(rows)}"
    )
    document.add_paragraph(
        "Times are provided in JST, UTC, and America/New_York."
    )

    display_fields = (
        "listing_id",
        "ended_at_new_york",
        "media_type",
        "disc_count",
        "bid_count",
        "start_price",
        "hammer_price_jpy",
        "tax_jpy",
        "gross_price_jpy",
        "seller",
        "title",
    )

    table = document.add_table(
        rows=1,
        cols=len(display_fields),
    )
    table.style = "Table Grid"

    for index, field in enumerate(
        display_fields,
    ):
        table.rows[0].cells[index].text = field

    for row in rows:
        cells = table.add_row().cells

        for index, field in enumerate(
            display_fields,
        ):
            value = row.get(field)
            cells[index].text = (
                "" if value is None else str(value)
            )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )
    document.save(output_path)

    return output_path


def main() -> int:
    args = parse_args()

    create_detail_table()

    candidates = load_candidates(
        force=args.force,
        limit=args.limit,
    )

    print()
    print("Buyee detail enrichment")
    print("-----------------------")
    print("Candidates:", len(candidates))

    changed = 0
    failed = 0

    if candidates:
        context = browser.context(
            args.profile
        )
        page = context.new_page()

        try:
            for index, row in enumerate(
                candidates,
                start=1,
            ):
                listing_id = row["listing_id"]

                print()
                print(
                    f"[{index}/{len(candidates)}] "
                    f"{listing_id}"
                )

                try:
                    detail = extract_detail(
                        page,
                        row,
                        args.wait_seconds,
                    )
                    save_detail(detail)
                    changed += 1

                    print(
                        "Ended NY :",
                        iso_in_zone(
                            detail["ended_at"],
                            NEW_YORK,
                        ),
                    )
                    print(
                        "Start JPY:",
                        detail["start_price"],
                    )
                    print(
                        "Bids     :",
                        detail["bid_count"],
                    )

                except Exception as exc:
                    failed += 1
                    print(
                        f"ERROR {listing_id}: {exc}"
                    )

                time.sleep(0.5)

        finally:
            page.close()

    rows = export_rows()

    csv_path = export_csv(
        rows,
        args.output_dir
        / "auctions_buyee_enriched.csv",
    )
    xlsx_path = export_xlsx(
        rows,
        args.output_dir
        / "auctions_buyee_enriched.xlsx",
    )
    docx_path = export_docx(
        rows,
        args.output_dir
        / "auctions_buyee_enriched.docx",
    )

    print()
    print("Enrichment summary")
    print("------------------")
    print("Changed :", changed)
    print("Failed  :", failed)
    print("Rows    :", len(rows))
    print("CSV     :", csv_path)
    print("XLSX    :", xlsx_path)
    print("DOCX    :", docx_path)

    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
