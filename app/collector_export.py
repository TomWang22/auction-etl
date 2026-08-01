"""Filtered and recent-ingestion exports for Collector Review."""

from __future__ import annotations

import json
import re
from datetime import date, datetime, timezone
from decimal import Decimal
from io import BytesIO
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
import streamlit as st

from app.collector_review_support import clean_text


ALL_COMBINED = "All media — one file"
ALL_SPLIT = "All media — ZIP by media"

EXPORT_FORMATS = (
    "CSV",
    "Excel",
    "JSON",
    "Markdown",
    "Word",
    "Plain text",
)

RECENT_FLAG_COLUMNS = (
    "_is_recent_addition",
    "is_recent_addition",
    "recent_addition",
    "recent_ingestion",
    "is_recent_ingestion",
)

INGESTION_BATCH_COLUMNS = (
    "ingestion_run_id",
    "ingestion_batch_id",
    "crawl_run_id",
    "crawl_job_id",
    "batch_id",
    "source_run_id",
)

INGESTION_DATE_COLUMNS = (
    "_audit_first_seen_at",
    "ingested_at",
    "added_at",
    "first_seen_at",
    "created_at",
)

MEDIA_COLUMNS = (
    "media_display",
    "effective_media_type",
    "manual_media_type",
    "auto_media_type",
    "media_type",
)

COLLECTOR_REPORT_FIELDS = (
    ("Marketplace", ("marketplace",)),
    ("Listing ID", ("listing_id",)),
    ("Seller", ("seller",)),
    ("Artist", ("artist_display", "artist")),
    ("Media", MEDIA_COLUMNS),
    (
        "Catalog / matrix",
        (
            "catalog_display",
            "pressing_token",
            "catalog_number",
        ),
    ),
    (
        "Pressing",
        (
            "effective_pressing_type",
            "manual_pressing_type",
            "auto_pressing_type",
            "edition",
        ),
    ),
    (
        "Condition media",
        (
            "effective_condition_media",
            "manual_condition_media",
            "condition_media",
        ),
    ),
    (
        "Condition cover",
        (
            "effective_condition_cover",
            "manual_condition_cover",
            "condition_cover",
        ),
    ),
    (
        "Starting price",
        (
            "starting_local",
            "start_price",
        ),
    ),
    (
        "Hammer before tax",
        (
            "hammer_local",
            "final_price",
        ),
    ),
    (
        "Tax",
        (
            "tax_local",
            "tax_amount",
        ),
    ),
    (
        "Total local",
        (
            "total_local",
            "gross_price",
            "current_price_gross",
        ),
    ),
    (
        "Total USD",
        (
            "total_usd",
            "gross_price_usd",
        ),
    ),
    (
        "Shipping",
        (
            "shipping_price",
            "shipping_usd",
        ),
    ),
    (
        "Bids",
        (
            "bid_count_display",
            "bid_count",
        ),
    ),
    (
        "Watch count",
        ("watch_count",),
    ),
    (
        "Closed",
        (
            "closing_display",
            "ended_at",
        ),
    ),
    (
        "Added",
        (
            "added_display",
            "_audit_first_seen_at",
            "created_at",
        ),
    ),
    (
        "Verdict",
        (
            "verdict_display",
            "manual_verdict",
        ),
    ),
    (
        "Collector notes",
        (
            "manual_collector_notes",
            "collector_notes",
        ),
    ),
    (
        "URL",
        ("auction_url",),
    ),
)


def _series(
    dataframe: pd.DataFrame,
    column: str,
    default: Any = "",
) -> pd.Series:
    """Return one Series even when duplicate column labels exist."""
    if column not in dataframe.columns:
        return pd.Series(
            default,
            index=dataframe.index,
            dtype="object",
        )

    selected = dataframe.loc[:, column]

    if isinstance(selected, pd.DataFrame):
        return (
            selected
            .bfill(axis=1)
            .iloc[:, 0]
        )

    return selected


def _truthy(value: Any) -> bool:
    """Interpret database, NumPy, and string booleans safely."""
    if value is None:
        return False

    try:
        missing = pd.isna(value)

        if (
            not hasattr(missing, "__len__")
            and bool(missing)
        ):
            return False
    except (TypeError, ValueError):
        pass

    if isinstance(value, bool):
        return value

    normalized = str(value).strip().lower()

    return normalized in {
        "1",
        "true",
        "t",
        "yes",
        "y",
        "recent",
    }


def _safe_value(value: Any) -> Any:
    """Convert one dataframe value into an export-safe scalar."""
    if value is None:
        return None

    if isinstance(value, Decimal):
        return str(value)

    if isinstance(
        value,
        (
            datetime,
            date,
            pd.Timestamp,
        ),
    ):
        try:
            if pd.isna(value):
                return None
        except (TypeError, ValueError):
            pass

        return value.isoformat()

    if isinstance(
        value,
        (
            dict,
            list,
            tuple,
            set,
        ),
    ):
        return json.dumps(
            value,
            ensure_ascii=False,
            default=str,
        )

    try:
        missing = pd.isna(value)

        if (
            not hasattr(missing, "__len__")
            and bool(missing)
        ):
            return None
    except (TypeError, ValueError):
        pass

    if hasattr(value, "item"):
        try:
            return value.item()
        except (TypeError, ValueError):
            pass

    return value


def _unique_column_names(
    columns: list[Any],
) -> list[str]:
    """Preserve every column while making duplicate names exportable."""
    counts: dict[str, int] = {}
    result: list[str] = []

    for position, raw_name in enumerate(
        columns,
        start=1,
    ):
        base_name = str(
            raw_name
            if raw_name not in {
                None,
                "",
            }
            else f"column_{position}"
        )

        counts[base_name] = (
            counts.get(
                base_name,
                0,
            )
            + 1
        )

        occurrence = counts[base_name]

        result.append(
            base_name
            if occurrence == 1
            else f"{base_name}__{occurrence}"
        )

    return result


def full_export_frame(
    dataframe: pd.DataFrame,
) -> pd.DataFrame:
    """Return every source column with JSON/CSV-safe values."""
    frame = dataframe.copy()

    frame.columns = _unique_column_names(
        list(frame.columns)
    )

    for column in frame.columns:
        frame[column] = frame[column].map(
            _safe_value
        )

    return frame


def _media_series(
    dataframe: pd.DataFrame,
) -> pd.Series:
    """Return the best available media classification."""
    for column in MEDIA_COLUMNS:
        if column not in dataframe.columns:
            continue

        values = (
            _series(
                dataframe,
                column,
            )
            .fillna("")
            .astype(str)
            .map(clean_text)
        )

        values = values.replace(
            "",
            "UNCLASSIFIED",
        )

        return values

    return pd.Series(
        "UNCLASSIFIED",
        index=dataframe.index,
        dtype="object",
    )


def recent_ingestion_frame(
    dataframe: pd.DataFrame,
) -> pd.DataFrame:
    """Select the established recent-ingestion set."""
    if dataframe.empty:
        return dataframe.copy()

    for column in RECENT_FLAG_COLUMNS:
        if column not in dataframe.columns:
            continue

        mask = _series(
            dataframe,
            column,
        ).map(_truthy)

        return (
            dataframe.loc[mask]
            .copy()
            .reset_index(drop=True)
        )

    for column in INGESTION_BATCH_COLUMNS:
        if column not in dataframe.columns:
            continue

        batches = _series(
            dataframe,
            column,
        )

        populated = batches[
            batches.notna()
            & batches.astype(str).str.strip().ne("")
        ]

        if populated.empty:
            continue

        latest_batch = populated.iloc[-1]

        return (
            dataframe.loc[
                batches.eq(latest_batch)
            ]
            .copy()
            .reset_index(drop=True)
        )

    for column in INGESTION_DATE_COLUMNS:
        if column not in dataframe.columns:
            continue

        timestamps = pd.to_datetime(
            _series(
                dataframe,
                column,
            ),
            errors="coerce",
            utc=True,
        )

        if timestamps.dropna().empty:
            continue

        latest_day = timestamps.max().date()

        return (
            dataframe.loc[
                timestamps.notna()
                & timestamps.dt.date.eq(
                    latest_day
                )
            ]
            .copy()
            .reset_index(drop=True)
        )

    return dataframe.iloc[0:0].copy()


def deduplicate_export_frame(
    dataframe: pd.DataFrame,
) -> pd.DataFrame:
    """Prefer native eBay over matching Gripsweat archive rows."""
    if dataframe.empty:
        return dataframe.copy()

    frame = dataframe.copy()

    marketplace = (
        _series(
            frame,
            "marketplace",
        )
        .fillna("")
        .astype(str)
        .str.strip()
        .str.lower()
    )

    listing_id = (
        _series(
            frame,
            "listing_id",
        )
        .fillna("")
        .astype(str)
        .str.strip()
    )

    source_marker = marketplace.copy()

    for source_column in (
        "_activity_source",
        "ingestion_source",
        "record_source",
        "data_source",
        "source",
    ):
        if source_column not in frame.columns:
            continue

        source_marker = (
            source_marker
            + " "
            + _series(
                frame,
                source_column,
            )
            .fillna("")
            .astype(str)
            .str.lower()
        )

    is_gripsweat = source_marker.str.contains(
        "gripsweat",
        regex=False,
    )

    frame["_export_family"] = marketplace.mask(
        is_gripsweat,
        "ebay",
    )

    frame["_export_priority"] = (
        is_gripsweat.astype(int)
    )

    frame["_export_order"] = range(
        len(frame)
    )

    frame["_export_listing_id"] = listing_id.mask(
        listing_id.eq(""),
        (
            "__missing__:"
            + frame[
                "_export_order"
            ].astype(str)
        ),
    )

    frame.sort_values(
        [
            "_export_priority",
            "_export_order",
        ],
        kind="stable",
        inplace=True,
    )

    frame.drop_duplicates(
        subset=[
            "_export_family",
            "_export_listing_id",
        ],
        keep="first",
        inplace=True,
    )

    frame.sort_values(
        "_export_order",
        kind="stable",
        inplace=True,
    )

    return (
        frame.drop(
            columns=[
                "_export_family",
                "_export_priority",
                "_export_order",
                "_export_listing_id",
            ],
        )
        .reset_index(drop=True)
    )


def _slug(value: str) -> str:
    """Return a filesystem-safe filename fragment."""
    result = re.sub(
        r"[^a-z0-9]+",
        "-",
        value.strip().lower(),
    ).strip("-")

    return result or "all"


def _row_value(
    row: pd.Series,
    candidates: tuple[str, ...],
) -> Any:
    """Return the first populated collector-report value."""
    for candidate in candidates:
        if candidate not in row.index:
            continue

        value = row[candidate]

        if isinstance(value, pd.Series):
            values = [
                item
                for item in value.tolist()
                if _safe_value(item) is not None
            ]

            if not values:
                continue

            value = values[0]

        normalized = _safe_value(value)

        if normalized not in {
            None,
            "",
        }:
            return normalized

    return None


def _csv_payload(
    dataframe: pd.DataFrame,
) -> tuple[bytes, str]:
    frame = full_export_frame(dataframe)

    return (
        frame.to_csv(
            index=False,
        ).encode("utf-8-sig"),
        "text/csv",
    )


def _json_payload(
    dataframe: pd.DataFrame,
) -> tuple[bytes, str]:
    frame = full_export_frame(dataframe)

    payload = json.dumps(
        frame.to_dict(
            orient="records"
        ),
        ensure_ascii=False,
        indent=2,
        default=str,
    ).encode("utf-8")

    return payload, "application/json"


def _text_payload(
    dataframe: pd.DataFrame,
) -> tuple[bytes, str]:
    frame = full_export_frame(dataframe)

    return (
        frame.to_csv(
            index=False,
            sep="\t",
        ).encode("utf-8"),
        "text/plain",
    )


def _excel_payload(
    dataframe: pd.DataFrame,
    title: str,
) -> tuple[bytes, str]:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    frame = full_export_frame(dataframe)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Listings"

    sheet.append(
        list(frame.columns)
    )

    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(
            fill_type="solid",
            fgColor="D9EAF7",
        )

    for values in frame.itertuples(
        index=False,
        name=None,
    ):
        sheet.append(
            list(values)
        )

    sheet.freeze_panes = "A2"

    if sheet.max_column > 0:
        sheet.auto_filter.ref = (
            sheet.dimensions
        )

    for column_index, column_name in enumerate(
        frame.columns,
        start=1,
    ):
        samples = [
            str(column_name),
            *[
                str(value)
                for value in frame[
                    column_name
                ].head(200)
                if value is not None
            ],
        ]

        width = min(
            70,
            max(
                12,
                max(
                    len(value)
                    for value in samples
                )
                + 2,
            ),
        )

        sheet.column_dimensions[
            get_column_letter(
                column_index
            )
        ].width = width

    summary = workbook.create_sheet(
        "Summary"
    )

    summary.append(
        [
            "Metric",
            "Value",
        ]
    )
    summary.append(
        [
            "Report",
            title,
        ]
    )
    summary.append(
        [
            "Generated UTC",
            datetime.now(
                timezone.utc
            ).isoformat(),
        ]
    )
    summary.append(
        [
            "Rows",
            len(frame),
        ]
    )
    summary.append(
        [
            "Columns",
            len(frame.columns),
        ]
    )

    for marketplace, count in sorted(
        _series(
            dataframe,
            "marketplace",
            "UNKNOWN",
        )
        .fillna("UNKNOWN")
        .astype(str)
        .value_counts()
        .items()
    ):
        summary.append(
            [
                f"Marketplace: {marketplace}",
                int(count),
            ]
        )

    for media_type, count in sorted(
        _media_series(
            dataframe
        )
        .value_counts()
        .items()
    ):
        summary.append(
            [
                f"Media: {media_type}",
                int(count),
            ]
        )

    for cell in summary[1]:
        cell.font = Font(bold=True)

    summary.column_dimensions["A"].width = 32
    summary.column_dimensions["B"].width = 60

    output = BytesIO()
    workbook.save(output)

    return (
        output.getvalue(),
        (
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )


def _markdown_payload(
    dataframe: pd.DataFrame,
    title: str,
) -> tuple[bytes, str]:
    media = _media_series(
        dataframe
    )

    lines = [
        f"# {title}",
        "",
        (
            "Generated: "
            + datetime.now(
                timezone.utc
            ).isoformat()
        ),
        "",
        f"Total listings: **{len(dataframe)}**",
        "",
    ]

    for media_type in sorted(
        media.unique()
    ):
        selected = dataframe.loc[
            media.eq(media_type)
        ]

        lines.extend(
            [
                f"## {media_type}",
                "",
                f"Listings: **{len(selected)}**",
                "",
            ]
        )

        for index, (_, row) in enumerate(
            selected.iterrows(),
            start=1,
        ):
            title_value = (
                _row_value(
                    row,
                    ("title",),
                )
                or "Untitled listing"
            )

            lines.extend(
                [
                    f"### {index}. {title_value}",
                    "",
                ]
            )

            for label, candidates in (
                COLLECTOR_REPORT_FIELDS
            ):
                value = _row_value(
                    row,
                    candidates,
                )

                if value in {
                    None,
                    "",
                }:
                    continue

                lines.append(
                    f"- **{label}:** {value}"
                )

            lines.append("")

    return (
        "\n".join(lines).encode(
            "utf-8"
        ),
        "text/markdown",
    )


def _word_payload(
    dataframe: pd.DataFrame,
    title: str,
) -> tuple[bytes, str]:
    from docx import Document

    document = Document()

    document.add_heading(
        title,
        level=0,
    )

    document.add_paragraph(
        f"Total listings: {len(dataframe)}"
    )

    document.add_paragraph(
        (
            "Generated UTC: "
            + datetime.now(
                timezone.utc
            ).isoformat()
        )
    )

    media = _media_series(
        dataframe
    )

    for media_type in sorted(
        media.unique()
    ):
        selected = dataframe.loc[
            media.eq(media_type)
        ]

        document.add_heading(
            f"{media_type} ({len(selected)})",
            level=1,
        )

        for index, (_, row) in enumerate(
            selected.iterrows(),
            start=1,
        ):
            title_value = (
                _row_value(
                    row,
                    ("title",),
                )
                or "Untitled listing"
            )

            document.add_heading(
                f"{index}. {title_value}",
                level=2,
            )

            table = document.add_table(
                rows=0,
                cols=2,
            )

            table.style = "Table Grid"

            for label, candidates in (
                COLLECTOR_REPORT_FIELDS
            ):
                value = _row_value(
                    row,
                    candidates,
                )

                if value in {
                    None,
                    "",
                }:
                    continue

                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = str(value)

    output = BytesIO()
    document.save(output)

    return (
        output.getvalue(),
        (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


def _single_payload(
    dataframe: pd.DataFrame,
    export_format: str,
    title: str,
) -> tuple[bytes, str, str]:
    """Build one file and return bytes, extension, and MIME type."""
    if export_format == "CSV":
        payload, mime_type = _csv_payload(
            dataframe
        )
        return payload, "csv", mime_type

    if export_format == "Excel":
        payload, mime_type = _excel_payload(
            dataframe,
            title,
        )
        return payload, "xlsx", mime_type

    if export_format == "JSON":
        payload, mime_type = _json_payload(
            dataframe
        )
        return payload, "json", mime_type

    if export_format == "Markdown":
        payload, mime_type = _markdown_payload(
            dataframe,
            title,
        )
        return payload, "md", mime_type

    if export_format == "Word":
        payload, mime_type = _word_payload(
            dataframe,
            title,
        )
        return payload, "docx", mime_type

    if export_format == "Plain text":
        payload, mime_type = _text_payload(
            dataframe
        )
        return payload, "txt", mime_type

    raise ValueError(
        f"Unsupported export format: {export_format}"
    )


def _filter_media(
    dataframe: pd.DataFrame,
    media_selection: str,
) -> pd.DataFrame:
    """Apply the export-specific media selector."""
    if media_selection in {
        ALL_COMBINED,
        ALL_SPLIT,
    }:
        return dataframe.copy()

    media = _media_series(
        dataframe
    )

    return (
        dataframe.loc[
            media.eq(media_selection)
        ]
        .copy()
        .reset_index(drop=True)
    )


def build_export_payload(
    dataframe: pd.DataFrame,
    export_format: str,
    media_selection: str,
    dataset_label: str,
) -> tuple[bytes, str, str, int]:
    """Build one file or a ZIP separated by media."""
    timestamp = datetime.now(
        timezone.utc
    ).strftime(
        "%Y%m%d-%H%M%S"
    )

    dataset_slug = _slug(
        dataset_label
    )

    if media_selection == ALL_SPLIT:
        media = _media_series(
            dataframe
        )

        archive_output = BytesIO()

        with ZipFile(
            archive_output,
            "w",
            compression=ZIP_DEFLATED,
        ) as archive:
            for media_type in sorted(
                media.unique()
            ):
                selected = (
                    dataframe.loc[
                        media.eq(media_type)
                    ]
                    .copy()
                    .reset_index(drop=True)
                )

                title = (
                    f"Auction Collector Review — "
                    f"{dataset_label} — {media_type}"
                )

                (
                    payload,
                    extension,
                    _,
                ) = _single_payload(
                    selected,
                    export_format,
                    title,
                )

                filename = (
                    "collector-review-"
                    f"{dataset_slug}-"
                    f"{_slug(media_type)}-"
                    f"{timestamp}."
                    f"{extension}"
                )

                archive.writestr(
                    filename,
                    payload,
                )

        return (
            archive_output.getvalue(),
            (
                "collector-review-"
                f"{dataset_slug}-"
                f"by-media-{timestamp}.zip"
            ),
            "application/zip",
            len(dataframe),
        )

    selected = _filter_media(
        dataframe,
        media_selection,
    )

    media_slug = (
        "all-media"
        if media_selection == ALL_COMBINED
        else _slug(media_selection)
    )

    title = (
        f"Auction Collector Review — "
        f"{dataset_label}"
    )

    (
        payload,
        extension,
        mime_type,
    ) = _single_payload(
        selected,
        export_format,
        title,
    )

    return (
        payload,
        (
            "collector-review-"
            f"{dataset_slug}-"
            f"{media_slug}-"
            f"{timestamp}."
            f"{extension}"
        ),
        mime_type,
        len(selected),
    )


def _frame_signature(
    dataframe: pd.DataFrame,
) -> str:
    """Return a stable signature for one exportable dataset."""
    if dataframe.empty:
        return "empty"

    signature_frame = pd.DataFrame(
        {
            "marketplace": _series(
                dataframe,
                "marketplace",
            ).astype(str),
            "listing_id": _series(
                dataframe,
                "listing_id",
            ).astype(str),
            "media": _media_series(
                dataframe
            ).astype(str),
            "total": _series(
                dataframe,
                "total_local",
            ).astype(str),
            "recent": _series(
                dataframe,
                "_is_recent_addition",
            ).astype(str),
        }
    )

    return str(
        int(
            pd.util.hash_pandas_object(
                signature_frame,
                index=False,
            ).sum()
        )
    )


def _prepare_export(
    dataframe: pd.DataFrame,
    export_format: str,
    media_selection: str,
    dataset_label: str,
) -> None:
    """Build and persist one download in Streamlit session state."""
    with st.spinner(
        f"Preparing {dataset_label.lower()} export…"
    ):
        (
            payload,
            filename,
            mime_type,
            row_count,
        ) = build_export_payload(
            dataframe,
            export_format,
            media_selection,
            dataset_label,
        )

    st.session_state[
        "_collector_export_result"
    ] = {
        "payload": payload,
        "filename": filename,
        "mime_type": mime_type,
        "row_count": row_count,
        "dataset_label": dataset_label,
        "format": export_format,
        "media": media_selection,
    }


def render_export_toolbar(
    filtered_dataframe: pd.DataFrame,
    all_dataframe: pd.DataFrame | None = None,
) -> None:
    """Render filtered, recent-filtered, and all-recent exports."""
    complete_dataframe = (
        all_dataframe
        if all_dataframe is not None
        else filtered_dataframe
    )

    with st.expander(
        "Export filtered or recent ingestion listings",
        expanded=False,
    ):
        st.caption(
            "CSV, Excel, JSON, and plain-text exports include every "
            "dataframe column. Markdown and Word exports are grouped "
            "by media and optimized for collector review."
        )

        (
            format_column,
            media_column,
            dedupe_column,
        ) = st.columns(
            [
                2,
                3,
                1,
            ]
        )

        with format_column:
            export_format = st.selectbox(
                "Format",
                EXPORT_FORMATS,
                key="collector_export_format",
            )

        media_source = pd.concat(
            [
                filtered_dataframe,
                complete_dataframe,
            ],
            ignore_index=True,
            sort=False,
        )

        media_values = tuple(
            sorted(
                set(
                    _media_series(
                        media_source
                    )
                    .dropna()
                    .astype(str)
                )
            )
        )

        with media_column:
            media_selection = st.selectbox(
                "Media",
                (
                    ALL_COMBINED,
                    ALL_SPLIT,
                    *media_values,
                ),
                key="collector_export_media",
            )

        with dedupe_column:
            deduplicate = st.checkbox(
                "Deduplicate",
                value=True,
                key="collector_export_deduplicate",
                help=(
                    "Prefer native eBay rows over matching "
                    "Gripsweat archive rows with the same listing ID."
                ),
            )

        filtered_export = (
            deduplicate_export_frame(
                filtered_dataframe
            )
            if deduplicate
            else filtered_dataframe.copy()
        )

        complete_export = (
            deduplicate_export_frame(
                complete_dataframe
            )
            if deduplicate
            else complete_dataframe.copy()
        )

        recent_filtered = recent_ingestion_frame(
            filtered_export
        )

        recent_all = recent_ingestion_frame(
            complete_export
        )

        filtered_selected = _filter_media(
            filtered_export,
            media_selection,
        )

        recent_filtered_selected = _filter_media(
            recent_filtered,
            media_selection,
        )

        recent_all_selected = _filter_media(
            recent_all,
            media_selection,
        )

        metrics = st.columns(3)

        metrics[0].metric(
            "Current filtered",
            len(filtered_selected),
        )

        metrics[1].metric(
            "Recent within filters",
            len(recent_filtered_selected),
        )

        metrics[2].metric(
            "All recent ingestion",
            len(recent_all_selected),
        )

        controls_signature = "|".join(
            (
                export_format,
                media_selection,
                str(deduplicate),
                _frame_signature(
                    filtered_export
                ),
                _frame_signature(
                    recent_filtered
                ),
                _frame_signature(
                    recent_all
                ),
            )
        )

        if (
            st.session_state.get(
                "_collector_export_controls"
            )
            != controls_signature
        ):
            st.session_state[
                "_collector_export_controls"
            ] = controls_signature

            st.session_state.pop(
                "_collector_export_result",
                None,
            )

        (
            filtered_button,
            recent_filtered_button,
            recent_all_button,
        ) = st.columns(3)

        with filtered_button:
            if st.button(
                "Prepare filtered export",
                type="primary",
                width="stretch",
                disabled=(
                    len(filtered_selected)
                    == 0
                ),
                key=(
                    "collector_prepare_"
                    "filtered_export"
                ),
            ):
                try:
                    _prepare_export(
                        filtered_export,
                        export_format,
                        media_selection,
                        "Current filtered results",
                    )
                except Exception as error:
                    st.error(
                        "Could not create filtered export: "
                        f"{error}"
                    )

        with recent_filtered_button:
            if st.button(
                "Prepare recent-in-filter export",
                width="stretch",
                disabled=(
                    len(
                        recent_filtered_selected
                    )
                    == 0
                ),
                key=(
                    "collector_prepare_"
                    "recent_filtered_export"
                ),
            ):
                try:
                    _prepare_export(
                        recent_filtered,
                        export_format,
                        media_selection,
                        "Recent ingestion within filters",
                    )
                except Exception as error:
                    st.error(
                        "Could not create recent filtered export: "
                        f"{error}"
                    )

        with recent_all_button:
            if st.button(
                "Prepare all recent ingestion",
                width="stretch",
                disabled=(
                    len(recent_all_selected)
                    == 0
                ),
                key=(
                    "collector_prepare_"
                    "all_recent_export"
                ),
            ):
                try:
                    _prepare_export(
                        recent_all,
                        export_format,
                        media_selection,
                        "All recent ingestion",
                    )
                except Exception as error:
                    st.error(
                        "Could not create all-recent export: "
                        f"{error}"
                    )

        result = st.session_state.get(
            "_collector_export_result"
        )

        if result:
            st.success(
                (
                    f"Prepared {result['row_count']:,} listings from "
                    f"{result['dataset_label']}."
                )
            )

            st.download_button(
                "Download prepared export",
                data=result["payload"],
                file_name=result["filename"],
                mime=result["mime_type"],
                width="stretch",
                key="collector_download_export",
            )
