from __future__ import annotations

import csv
import json
from collections import Counter
from datetime import date, datetime, timezone
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.orm import Session

from auction_etl.models.warehouse import Auction
from auction_etl.services.fx import FxQuote, latest_rate


EXPORT_COLUMNS = (
    "id",
    "marketplace",
    "listing_id",
    "auction_url",
    "seller",
    "artist",
    "title",
    "media_type",
    "disc_count",
    "edition",
    "catalog_number",
    "condition_media",
    "condition_cover",
    "bulk_lot",
    "bid_count",
    "watch_count",
    "start_price",
    "hammer_price_local",
    "tax_rate",
    "tax_amount_local",
    "gross_price_local",
    "price_includes_tax",
    "currency",
    "fx_rate_to_usd",
    "fx_rate_date",
    "final_price_usd",
    "tax_usd",
    "total_usd",
    "shipping_price",
    "shipping_usd",
    "landed_usd",
    "ended_at",
    "created_at",
)


def _decimal(
    value: Any,
) -> Decimal | None:
    if value is None:
        return None

    return Decimal(str(value))


def _money(
    value: Decimal | int | float | str | None,
) -> Decimal | None:
    if value is None:
        return None

    decimal_value = (
        value
        if isinstance(value, Decimal)
        else Decimal(str(value))
    )

    return decimal_value.quantize(
        Decimal("0.01"),
        rounding=ROUND_HALF_UP,
    )


def _serialize(value: Any) -> Any:
    if isinstance(value, Decimal):
        return str(value)

    if isinstance(value, (datetime, date)):
        return value.isoformat()

    return value


def _quotes(
    currencies: set[str],
    *,
    live_fx: bool,
    jpy_usd_rate: Decimal | None,
) -> dict[str, FxQuote]:
    today = date.today()

    quotes = {
        "USD": FxQuote(
            base="USD",
            quote="USD",
            rate=Decimal("1"),
            rate_date=today,
            source="identity",
        )
    }

    for currency in currencies:
        if currency == "USD":
            continue

        if (
            currency == "JPY"
            and jpy_usd_rate is not None
        ):
            quotes[currency] = FxQuote(
                base="JPY",
                quote="USD",
                rate=jpy_usd_rate,
                rate_date=today,
                source="manual override",
            )
            continue

        if not live_fx:
            raise RuntimeError(
                f"No manual USD rate supplied for {currency}."
            )

        quotes[currency] = latest_rate(
            currency,
            "USD",
        )

    return quotes


def load_rows(
    session: Session,
    marketplace: str | None = None,
    *,
    live_fx: bool = True,
    jpy_usd_rate: Decimal | None = None,
) -> list[dict[str, Any]]:
    statement = select(Auction).order_by(
        Auction.marketplace,
        Auction.id,
    )

    if marketplace:
        statement = statement.where(
            Auction.marketplace
            == marketplace
        )

    auctions = session.scalars(
        statement
    ).all()

    currencies = {
        (auction.currency or "USD").upper()
        for auction in auctions
    }

    quotes = _quotes(
        currencies,
        live_fx=live_fx,
        jpy_usd_rate=jpy_usd_rate,
    )

    rows: list[dict[str, Any]] = []

    for auction in auctions:
        currency = (
            auction.currency or "USD"
        ).upper()

        quote = quotes[currency]

        hammer_price = _decimal(
            auction.final_price
        )
        tax_amount = _decimal(
            auction.tax_amount
        )
        gross_price = _decimal(
            auction.gross_price
        )

        if (
            gross_price is None
            and hammer_price is not None
        ):
            gross_price = (
                hammer_price
                + (
                    tax_amount
                    or Decimal("0")
                )
            )

        tax_for_math = (
            tax_amount
            or Decimal("0")
        )

        shipping_price = (
            _decimal(
                auction.shipping_price
            )
            or Decimal("0")
        )

        final_price_usd = (
            _money(
                hammer_price * quote.rate
            )
            if hammer_price is not None
            else None
        )

        tax_usd = (
            _money(
                tax_amount * quote.rate
            )
            if tax_amount is not None
            else None
        )

        total_usd = (
            _money(
                gross_price * quote.rate
            )
            if gross_price is not None
            else None
        )

        shipping_usd = _money(
            shipping_price * quote.rate
        )

        landed_usd = (
            _money(
                (
                    gross_price
                    + shipping_price
                )
                * quote.rate
            )
            if gross_price is not None
            else None
        )

        rows.append(
            {
                "id": auction.id,
                "marketplace": (
                    auction.marketplace
                ),
                "listing_id": (
                    auction.listing_id
                ),
                "auction_url": (
                    auction.auction_url
                ),
                "seller": auction.seller,
                "artist": auction.artist,
                "title": auction.title,
                "media_type": (
                    auction.media_type
                ),
                "disc_count": (
                    auction.disc_count
                ),
                "edition": auction.edition,
                "catalog_number": (
                    auction.catalog_number
                ),
                "condition_media": (
                    auction.condition_media
                ),
                "condition_cover": (
                    auction.condition_cover
                ),
                "bulk_lot": auction.bulk_lot,
                "bid_count": (
                    auction.bid_count
                ),
                "watch_count": (
                    auction.watch_count
                ),
                "start_price": (
                    auction.start_price
                ),
                "hammer_price_local": hammer_price,
                "tax_rate": auction.tax_rate,
                "tax_amount_local": tax_amount,
                "gross_price_local": gross_price,
                "price_includes_tax": (
                    auction.price_includes_tax
                ),
                "currency": currency,
                "fx_rate_to_usd": (
                    quote.rate
                ),
                "fx_rate_date": (
                    quote.rate_date
                ),
                "final_price_usd": (
                    final_price_usd
                ),
                "tax_usd": tax_usd,
                "total_usd": total_usd,
                "shipping_price": (
                    shipping_price
                ),
                "shipping_usd": (
                    shipping_usd
                ),
                "landed_usd": landed_usd,
                "ended_at": (
                    auction.ended_at
                ),
                "created_at": (
                    auction.created_at
                ),
                "_fx_source": quote.source,
            }
        )

    return rows


def ensure_output_dir(
    output_dir: str | Path,
) -> Path:
    path = Path(output_dir)
    path.mkdir(
        parents=True,
        exist_ok=True,
    )
    return path


def export_csv(
    rows: list[dict[str, Any]],
    output_path: str | Path,
) -> Path:
    path = Path(output_path)

    with path.open(
        "w",
        encoding="utf-8-sig",
        newline="",
    ) as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                EXPORT_COLUMNS
            ),
            extrasaction="ignore",
        )

        writer.writeheader()

        for row in rows:
            writer.writerow(
                {
                    key: _serialize(
                        row.get(key)
                    )
                    for key in EXPORT_COLUMNS
                }
            )

    return path


def export_json(
    rows: list[dict[str, Any]],
    output_path: str | Path,
) -> Path:
    path = Path(output_path)

    payload = [
        {
            key: _serialize(
                value
            )
            for key, value
            in row.items()
            if not key.startswith("_")
        }
        for row in rows
    ]

    path.write_text(
        json.dumps(
            payload,
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    return path


def export_markdown(
    rows: list[dict[str, Any]],
    output_path: str | Path,
    title: str = "Auction Listings",
) -> Path:
    path = Path(output_path)

    lines = [
        f"# {title}",
        "",
        f"Generated: {datetime.now(timezone.utc).isoformat()}",
        "",
        f"Total listings: **{len(rows)}**",
        "",
        "| Marketplace | Listings |",
        "|---|---:|",
    ]

    for marketplace, count in sorted(
        Counter(
            row["marketplace"]
            for row in rows
        ).items()
    ):
        lines.append(
            f"| {marketplace} | {count} |"
        )

    lines.extend(
        [
            "",
            "## Listings",
            "",
            "| Marketplace | Ended | Title | Original | Tax | Total USD | Seller |",
            "|---|---|---|---:|---:|---:|---|",
        ]
    )

    for row in rows:
        title_text = str(
            row["title"] or ""
        ).replace("|", r"\|")

        lines.append(
            "| "
            f"{row['marketplace']} | "
            f"{row['ended_at'] or ''} | "
            f"{title_text} | "
            f"{row['hammer_price_local'] or ''} "
            f"{row['currency']} | "
            f"{row['tax_amount_local'] or 0} | "
            f"{row['total_usd'] or ''} | "
            f"{row['seller'] or ''} |"
        )

    path.write_text(
        "\n".join(lines),
        encoding="utf-8",
    )

    return path


def export_xlsx(
    rows: list[dict[str, Any]],
    output_path: str | Path,
) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.table import Table
    from openpyxl.worksheet.table import TableStyleInfo

    path = Path(output_path)
    workbook = Workbook()

    sheet = workbook.active
    sheet.title = "Auctions"

    sheet.append(
        list(EXPORT_COLUMNS)
    )

    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(
            fill_type="solid",
            fgColor="D9EAF7",
        )

    for row in rows:
        sheet.append(
            [
                _serialize(
                    row.get(column)
                )
                for column in EXPORT_COLUMNS
            ]
        )

    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = (
        sheet.dimensions
    )

    if rows:
        table = Table(
            displayName="AuctionWarehouse",
            ref=sheet.dimensions,
        )
        table.tableStyleInfo = (
            TableStyleInfo(
                name="TableStyleMedium2",
                showFirstColumn=False,
                showLastColumn=False,
                showRowStripes=True,
                showColumnStripes=False,
            )
        )
        sheet.add_table(table)

    money_columns = {
        "start_price",
        "hammer_price_local",
        "tax_amount_local",
        "gross_price_local",
        "final_price_usd",
        "tax_usd",
        "total_usd",
        "shipping_price",
        "shipping_usd",
        "landed_usd",
    }

    rate_columns = {
        "fx_rate_to_usd",
    }

    for index, column in enumerate(
        EXPORT_COLUMNS,
        start=1,
    ):
        width = 18

        if column == "title":
            width = 65
        elif column == "auction_url":
            width = 55
        elif column in {
            "artist",
            "seller",
        }:
            width = 28

        sheet.column_dimensions[
            get_column_letter(index)
        ].width = width

        if column in money_columns:
            for cell in sheet[
                get_column_letter(index)
            ][1:]:
                cell.number_format = (
                    '#,##0.00'
                )

        if column in rate_columns:
            for cell in sheet[
                get_column_letter(index)
            ][1:]:
                cell.number_format = (
                    '0.00000000'
                )

    settings = workbook.create_sheet(
        "FX Settings"
    )

    settings.append(
        ["Setting", "Value"]
    )
    settings.append(
        [
            "Generated UTC",
            datetime.now(
                timezone.utc
            ).isoformat(),
        ]
    )

    seen: set[
        tuple[str, str, str]
    ] = set()

    for row in rows:
        key = (
            row["currency"],
            str(row["fx_rate_to_usd"]),
            str(row["fx_rate_date"]),
        )

        if key in seen:
            continue

        seen.add(key)

        settings.append(
            [
                f"{row['currency']} → USD",
                row["fx_rate_to_usd"],
            ]
        )
        settings.append(
            [
                f"{row['currency']} rate date",
                str(row["fx_rate_date"]),
            ]
        )
        settings.append(
            [
                f"{row['currency']} source",
                row["_fx_source"],
            ]
        )

    for cell in settings[1]:
        cell.font = Font(bold=True)

    settings.column_dimensions["A"].width = 28
    settings.column_dimensions["B"].width = 42

    summary = workbook.create_sheet(
        "Summary"
    )

    summary.append(
        ["Metric", "Value"]
    )
    summary.append(
        ["Total auctions", len(rows)]
    )

    total_usd = sum(
        (
            row.get("total_usd")
            or Decimal("0")
        )
        for row in rows
    )

    landed_usd = sum(
        (
            row.get("landed_usd")
            or Decimal("0")
        )
        for row in rows
    )

    summary.append(
        [
            "Auction totals USD",
            _money(total_usd),
        ]
    )
    summary.append(
        [
            "Landed totals USD",
            _money(landed_usd),
        ]
    )

    for marketplace, count in sorted(
        Counter(
            row["marketplace"]
            for row in rows
        ).items()
    ):
        summary.append(
            [
                f"{marketplace} auctions",
                count,
            ]
        )

    for cell in summary[1]:
        cell.font = Font(bold=True)

    summary.column_dimensions["A"].width = 28
    summary.column_dimensions["B"].width = 22

    workbook.save(path)
    return path


def export_docx(
    rows: list[dict[str, Any]],
    output_path: str | Path,
    title: str = "Auction Listings",
) -> Path:
    from docx import Document

    path = Path(output_path)
    document = Document()

    document.add_heading(
        title,
        level=0,
    )

    document.add_paragraph(
        f"Total auctions: {len(rows)}"
    )

    document.add_paragraph(
        "USD values use the FX rate shown "
        "for each auction."
    )

    fields = (
        ("Marketplace", "marketplace"),
        ("Listing ID", "listing_id"),
        ("Artist", "artist"),
        ("Media type", "media_type"),
        ("Disc count", "disc_count"),
        ("Catalog number", "catalog_number"),
        ("Hammer price", "hammer_price_local"),
        ("Tax rate", "tax_rate"),
        ("Tax amount", "tax_amount_local"),
        ("Gross price", "gross_price_local"),
        ("Price includes tax", "price_includes_tax"),
        ("Original currency", "currency"),
        ("FX rate to USD", "fx_rate_to_usd"),
        ("FX rate date", "fx_rate_date"),
        ("Total USD", "total_usd"),
        ("Shipping USD", "shipping_usd"),
        ("Landed USD", "landed_usd"),
        ("Seller", "seller"),
        ("Ended at", "ended_at"),
        ("URL", "auction_url"),
    )

    for index, row in enumerate(
        rows,
        start=1,
    ):
        document.add_heading(
            f"{index}. "
            f"{row['title']}",
            level=1,
        )

        table = document.add_table(
            rows=0,
            cols=2,
        )
        table.style = "Table Grid"

        for label, key in fields:
            value = row.get(key)

            if value in {
                None,
                "",
            }:
                continue

            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value)

    document.save(path)
    return path
